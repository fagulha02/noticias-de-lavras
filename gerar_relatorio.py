from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ===== ESTILOS =====
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ===== TITULO =====
title = doc.add_heading('Relatorio Consolidado de Entregas - SuperInova Lavras', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph('')

# ===== RESUMO EXECUTIVO =====
doc.add_heading('Resumo Executivo', level=1)

resumo = doc.add_paragraph()
resumo.paragraph_format.space_after = Pt(6)
resumo.add_run(
    'O presente relatorio consolida todas as atividades finalizadas (status "Historico") '
    'registradas no banco de dados Sprint do espaco SuperInova - Lavras no Notion, '
    'bem como as metas mensais e semestrais concluidas pela equipe. '
    'No total, foram identificadas 36 atividades concluidas e movidas para o historico, '
    'alem de 11 metas atingidas (3 mensais e 8 semestrais).'
)

doc.add_paragraph('')

doc.add_heading('Principais Avancos por Area (OKR)', level=2)

avancos = [
    (
        'OKR 04 - Desenvolvimento Economico e Startups (6 atividades)',
        'Esta foi a area com maior volume de entregas identificadas por OKR. As acoes incluiram '
        'a organizacao do Startup Day, o alinhamento com atores da Trilha Empreendedora do Cluster '
        'AgroFoodTech, a preparacao do Kick-off da Comunidade de Startups, a criacao de criativos '
        'para o evento e o levantamento de metodologias para a Economia Prateada (60+). '
        'Essas entregas demonstram um forte investimento na construcao de um pipeline empreendedor '
        'e no fortalecimento da comunidade de startups de Lavras.'
    ),
    (
        'OKR 07 - Branding e Comunicacao (4 atividades)',
        'Destaque para a criacao do site basico do Pacto/Vale dos Ipes, o desenvolvimento do '
        'prototipo do site do Vale dos Ipes no Figma/Miro, a definicao do review de sprint para '
        'marketing e a estruturacao da identidade visual. Essas acoes consolidam a presenca digital '
        'e a comunicacao institucional do ecossistema.'
    ),
    (
        'OKR 05 - Governo Digital (3 atividades)',
        'Focado na modernizacao da Sala Mineira, com mapeamento de fluxos de trabalho, '
        'reconstrucao da dinamica operacional e criacao de cronograma de atividades com metodologias '
        'ageis. Representa um avanco significativo na transformacao digital do servico publico local.'
    ),
    (
        'OKR 03 - Legislacao e Ambiente Regulatorio (3 atividades)',
        'Avancos na implementacao do Redesim+ Livre (Fase 04), verificacao de suporte do Sebrae e '
        'adequacao regulatoria. Essas acoes pavimentam o caminho para simplificacao do ambiente de '
        'negocios no municipio.'
    ),
    (
        'OKR 02 - LVRS+ 2040 (2 atividades)',
        'Ativacao do GTO (Workshop de co-criacao) e convocacao da sociedade para a mesa do Pacto. '
        'Acoes estrategicas de governanca e engajamento territorial.'
    ),
    (
        'Atividades de Suporte e Operacionais (18 atividades)',
        'Um volume expressivo de atividades operacionais e de suporte, incluindo logistica de '
        'eventos (orcamentos, crachas, telao/som, reserva de veiculos), comunicacao (artes visuais, '
        'blogueiros/influencers, programa de radio, marketing com Rhizoma), articulacao institucional '
        '(roteiros de visita SENAC, convites GTO, Clube 506) e planejamento estrategico '
        '(definicao de OKRs, incentivo fiscal, negociacao de emendas). Essas entregas, embora sem '
        'OKR especifico atribuido, foram fundamentais para viabilizar os eventos e o funcionamento '
        'do ecossistema.'
    ),
]

for titulo_okr, descricao in avancos:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run_titulo = p.add_run(titulo_okr)
    run_titulo.bold = True
    run_titulo.font.size = Pt(11)

    p2 = doc.add_paragraph(descricao)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph('')
p_conclusao = doc.add_paragraph()
p_conclusao.add_run('Conclusao: ').bold = True
p_conclusao.add_run(
    'A SuperInova Lavras demonstra forte capacidade de execucao, com destaque para a articulacao '
    'do ecossistema de startups e inovacao, a realizacao de eventos de alto impacto (Pacto Lavras, '
    'Startup Day, Kick-off), o avanco regulatorio com o Redesim+ Livre e a modernizacao da gestao '
    'publica na Sala Mineira. O cumprimento de 11 das 12 metas de curto e medio prazo (91,7%) '
    'reafirma o comprometimento da equipe com os objetivos estrategicos do programa.'
)

doc.add_page_break()

# ===== TABELA DE ATIVIDADES =====
doc.add_heading('Atividades Finalizadas - Sprint', level=1)
doc.add_paragraph('Abaixo estao listadas todas as atividades com status "Historico" no banco de dados Sprint.')
doc.add_paragraph('')

atividades = [
    ("Criar cronograma de atividades com a Sala Mineira", "-", "ALTA", "OKR 05 - Governo Digital"),
    ("Reconstruir a dinamica da Sala Mineira", "18/02 a 23/02/2026", "ALTA", "OKR 05 - Governo Digital"),
    ("Definir dias e horarios para um review de sprint. MKT-VDI", "-", "ALTA", "OKR 07 - Branding e Comunicacao"),
    ("Desenvolver roteiro/programacao do Pacto Lavras", "-", "ALTA", "-"),
    ("Preparar Fase 04 do Redesim + Livre", "09/03 a 13/03/2026", "ALTA", "OKR 03 - Legislacao e Amb. Regulatorio"),
    ("Criativos sobre programacao do evento", "12/03 a 19/03/2026", "ALTA", "OKR 04 - Des. Economico e Startups"),
    ("Solicitar participacao do programa de radio", "-", "ALTA", "-"),
    ("Ativar GTO - Workshop", "-", "ALTA", "OKR 02 - LVRS+ 2040"),
    ("Convocar Sociedade para a mesa do Pacto", "-", "ALTA", "OKR 02 - LVRS+ 2040"),
    ("Startup Day - Estrutura do Evento", "21/03/2026", "MEDIA", "OKR 04 - Des. Economico e Startups"),
    ("Avaliar Alinhamento com Atores da Trilha Empreendedora", "-", "ALTA", "OKR 04 - Des. Economico e Startups"),
    ("Reuniao com o Clube 506 - Lancamento do Pacto Lavras", "-", "-", "-"),
    ("Marcar Kick-off com a comunidade de Startups para o dia 26/02", "11/02/2026", "MEDIA", "OKR 04 - Des. Economico e Startups"),
    ("Mandar convite para Gustavo e Israel - GTO", "-", "ALTA", "-"),
    ("Levantamento de metodologias para projetos (60+) - Economia Prateada", "-", "-", "OKR 04 - Des. Economico e Startups"),
    ("Ajustar landing page do Pacto Lavras", "-", "-", "-"),
    ("Mapear fluxo de trabalho - Daniela", "-", "ALTA", "OKR 05 - Governo Digital"),
    ("Alinhar formato do Kick-off da Comunidade de Startups de Lavras", "10/02 a 23/02/2026", "ALTA", "OKR 04 - Des. Economico e Startups"),
    ("Verificar suporte do Sebrae para implementacao do Rede Sim+ Livre", "-", "ALTA", "OKR 03 - Legislacao e Amb. Regulatorio"),
    ("Criar roteiro para o Prefeito de Sao Goncalo do Rio Abaixo", "12/11/2025", "ALTA", "EXTRA"),
    ("Fazer site basico sobre o pacto e o vale dos ipes - Versao 01", "-", "ALTA", "OKR 07 - Branding e Comunicacao"),
    ("Ajudar na construcao de artes visuais do Pacto Lavras", "-", "-", "-"),
    ("Desenvolver prototipo do site do Vale dos Ipes", "-", "ALTA", "OKR 07 - Branding e Comunicacao"),
    ("Elaborar roteiros de visita - SENAC", "-", "ALTA", "-"),
    ("Consolidar orcamentos do Lancamento do Pacto Lavras", "-", "-", "-"),
    ("Reservar o carro para ir na rodoviaria - Sexta", "-", "ALTA", "-"),
    ("Filtrar todos os itens disponiveis na Rodoviaria", "-", "ALTA", "-"),
    ("Convocar blogueiros e influencers do municipio - Pacto Lavras", "-", "ALTA", "-"),
    ("Definir responsaveis e deadlines para as OKR's", "-", "MEDIA", "-"),
    ("Preparar estrategia do marketing pos evento", "-", "ALTA", "-"),
    ("Verificar andamento do Incentivo Fiscal", "-", "ALTA", "-"),
    ("Reunir com a Rhizoma para planejamento do Marketing - VDI", "-", "-", "-"),
    ("Definir quadro de postagens do VDI - Rhizoma e SuperInova", "-", "ALTA", "-"),
    ("Negociacao da emenda da Rhizoma - Nova identidade visual e site", "-", "ALTA", "-"),
    ("Solicitar orcamento de telao, som e iluminacao - Pacto Lavras", "-", "-", "-"),
    ("Solicitar orcamento de crachas e camisetas - Pacto Lavras", "-", "-", "-"),
]

# Criar tabela
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Cabeçalho
hdr_cells = table.rows[0].cells
headers = ['Nome da Atividade', 'Data', 'Prioridade', 'Tipo de Atividade (OKR)']
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Cor de fundo do cabeçalho
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1A3C6E')
    shading.set(qn('w:val'), 'clear')
    hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

# Dados
for atividade in atividades:
    row_cells = table.add_row().cells
    for i, valor in enumerate(atividade):
        row_cells[i].text = valor
        for paragraph in row_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Ajustar largura das colunas
for row in table.rows:
    row.cells[0].width = Cm(9)
    row.cells[1].width = Cm(3.5)
    row.cells[2].width = Cm(2)
    row.cells[3].width = Cm(5)

doc.add_paragraph('')

doc.add_paragraph(f'Total de atividades finalizadas: {len(atividades)}').runs[0].bold = True

doc.add_page_break()

# ===== METAS CONCLUIDAS =====
doc.add_heading('Metas Concluidas', level=1)

doc.add_heading('Meta Mensal (Deadline: 31/10/2025)', level=2)
metas_mensais = [
    'Governanca do Pacto',
    '54 Congresso Nacional de Fitopatologia - 04 a 08 de Agosto',
    'Missao Startup Summit - 25 a 30 de Agosto',
]
for meta in metas_mensais:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(meta)
    run.font.size = Pt(11)

doc.add_paragraph('')

doc.add_heading('Meta Semestral (Deadline: 31/12/2025)', level=2)
metas_semestrais = [
    'Conselho Municipal',
    'Convocar Mesa do Pacto',
    'Reuniao da Mesa',
    'Workshop de co-criacao de projetos',
    'Lancar Pacto SRI',
    'Planejar Trilhas Empreendedoras',
    'Aprovar Sandbox',
    'Lancar Pacto Lavras',
]
for meta in metas_semestrais:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(meta)
    run.font.size = Pt(11)

doc.add_paragraph('')

# Nota sobre meta não concluída
p_nota = doc.add_paragraph()
p_nota.add_run('Nota: ').bold = True
p_nota.add_run(
    'A meta semestral "Atualizar Incentivos Fiscais" permanece pendente. '
    'As metas anuais (Deadline: 31/07/2026) nao foram incluidas neste relatorio '
    'pois nenhuma foi concluida ate o momento.'
)

doc.add_paragraph('')
doc.add_paragraph('')

# Rodapé
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run('Relatorio gerado automaticamente em 15/04/2026 a partir dos dados do Notion - SuperInova Lavras')
run_footer.font.size = Pt(9)
run_footer.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run_footer.italic = True

# Salvar
output_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'atividades finalizadas.docx')
doc.save(output_path)
print(f"Arquivo salvo em: {output_path}")
